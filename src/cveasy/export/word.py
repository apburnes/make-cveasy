"""Word document export functionality."""

from pathlib import Path
from typing import TYPE_CHECKING
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

from cveasy.export.styles import DEFAULT_STYLE

if TYPE_CHECKING:
    from cveasy.export.styles import ResumeStyle


def markdown_to_docx_paragraphs(
    doc: Document,
    text: str,
    resume_style: "ResumeStyle" = DEFAULT_STYLE,  # noqa: F821
):
    """Convert markdown text to docx paragraphs."""
    lines = text.split("\n")
    previous_was_header = False

    for line in lines:
        line = line.strip()

        # Skip markdown dividers
        if line == "---" or line.startswith("---"):
            continue

        if not line:
            continue

        # Headers
        if line.startswith("# "):
            # Add spacing before header if not first element
            if previous_was_header:
                # Add a small spacer paragraph
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(resume_style.spacing_heading_before * 72)
            p = doc.add_heading(line[2:], level=1)
            p.paragraph_format.space_after = Pt(resume_style.spacing_heading_after * 72)
            # Set font size and weight (H1: bold/700)
            for run in p.runs:
                run.font.size = Pt(resume_style.font_size_h1)
                run.bold = True  # Weight 700
            previous_was_header = True
            continue
        elif line.startswith("## "):
            # Add spacing before header if not first element
            if previous_was_header:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(resume_style.spacing_heading_before * 72)
            p = doc.add_heading(line[3:], level=2)
            p.paragraph_format.space_after = Pt(resume_style.spacing_heading_after * 72)
            # Set font size and weight (H2: semi-bold/600)
            for run in p.runs:
                run.font.size = Pt(resume_style.font_size_h2)
                # python-docx doesn't support semi-bold directly, use bold as approximation
                run.bold = True
            previous_was_header = True
            continue
        elif line.startswith("### "):
            # Add spacing before subheader
            if previous_was_header:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(resume_style.spacing_heading_before * 72)
            p = doc.add_heading(line[4:], level=3)
            p.paragraph_format.space_after = Pt(resume_style.spacing_heading_after * 72)
            # Set font size and weight (H3: medium/500)
            for run in p.runs:
                run.font.size = Pt(resume_style.font_size_h3)
                # Medium weight - keep regular (not bold)
                run.bold = False
            previous_was_header = True
            continue

        previous_was_header = False

        # Lists
        if line.startswith("- ") or line.startswith("* "):
            list_text = line[2:].strip()
            # Process inline formatting for list items
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(resume_style.spacing_list_item * 72)
            p.paragraph_format.left_indent = Inches(resume_style.list_indent)

            # Process inline formatting
            parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))", list_text)

            for part in parts:
                if not part:
                    continue

                # Bold
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(resume_style.font_size_body)
                # Italic
                elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                    run = p.add_run(part[1:-1])
                    run.italic = True
                    run.font.size = Pt(resume_style.font_size_body)
                # Links
                elif part.startswith("[") and "](" in part:
                    link_match = re.match(r"\[(.+?)\]\((.+?)\)", part)
                    if link_match:
                        link_text = link_match.group(1)
                        link_url = link_match.group(2)
                        run = p.add_run(link_text)
                        run.font.size = Pt(resume_style.font_size_body)
                        if resume_style.link_preserve_url:
                            # Add URL text after link text
                            run = p.add_run(f" ({link_url})")
                            run.font.size = Pt(resume_style.font_size_body)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(resume_style.font_size_body)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(resume_style.spacing_paragraph * 72)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Process inline formatting
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))", line)

        for part in parts:
            if not part:
                continue

            # Bold
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(resume_style.font_size_body)
            # Italic
            elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                run = p.add_run(part[1:-1])
                run.italic = True
                run.font.size = Pt(resume_style.font_size_body)
            # Links
            elif part.startswith("[") and "](" in part:
                link_match = re.match(r"\[(.+?)\]\((.+?)\)", part)
                if link_match:
                    link_text = link_match.group(1)
                    link_url = link_match.group(2)
                    run = p.add_run(link_text)
                    run.font.size = Pt(resume_style.font_size_body)
                    if resume_style.link_preserve_url:
                        # Add hyperlink (simplified - just add URL text)
                        run = p.add_run(f" ({link_url})")
                        run.font.size = Pt(resume_style.font_size_body)
            else:
                run = p.add_run(part)
                run.font.size = Pt(resume_style.font_size_body)


def export_to_word(
    resume_text: str,
    output_path: Path,
    resume_style: "ResumeStyle" = DEFAULT_STYLE,  # noqa: F821
) -> Path:
    """
    Export resume to Word document.

    Args:
        resume_text: Resume content in markdown
        output_path: Output file path
        resume_style: Style configuration to use (defaults to DEFAULT_STYLE)

    Returns:
        Path to created Word document
    """
    doc = Document()

    # Set margins using shared styles
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(resume_style.margin_top)
        section.bottom_margin = Inches(resume_style.margin_bottom)
        section.left_margin = Inches(resume_style.margin_left)
        section.right_margin = Inches(resume_style.margin_right)

    # Convert markdown to docx
    markdown_to_docx_paragraphs(doc, resume_text, resume_style)

    # Save document
    doc.save(str(output_path))

    return output_path
